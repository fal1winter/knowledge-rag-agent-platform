# Python RAG服务 - 文档解析器实现
# 路径: /home/sun/javabackend/material-rag-service/document_parser.py

import requests
from typing import List, Dict
import PyPDF2
from docx import Document
import markdown
from bs4 import BeautifulSoup
import logging
import io
import re

logger = logging.getLogger(__name__)

def _split_text(text: str, chunk_size=500, chunk_overlap=50) -> List[str]:
    """简单的递归文本分块，替代langchain的RecursiveCharacterTextSplitter"""
    separators = ["\n\n", "\n", "。", "！", "？", ".", "!", "?", " "]
    chunks = []
    _recursive_split(text, separators, chunk_size, chunk_overlap, chunks)
    return chunks

def _recursive_split(text, separators, chunk_size, overlap, result):
    if len(text) <= chunk_size:
        if text.strip():
            result.append(text.strip())
        return
    sep = ""
    for s in separators:
        if s in text:
            sep = s
            break
    if not sep:
        # 硬切
        for i in range(0, len(text), chunk_size - overlap):
            piece = text[i:i + chunk_size]
            if piece.strip():
                result.append(piece.strip())
        return
    parts = text.split(sep)
    current = ""
    for part in parts:
        candidate = (current + sep + part) if current else part
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            if current.strip():
                result.append(current.strip())
            if len(part) > chunk_size:
                remaining_seps = separators[separators.index(sep)+1:] if sep in separators else []
                _recursive_split(part, remaining_seps, chunk_size, overlap, result)
                current = ""
            else:
                current = part
    if current.strip():
        result.append(current.strip())

class DocumentParser:
    def __init__(self):
        self.timeout = 60

    def parse(self, file_url: str, file_type: str) -> List[Dict]:
        """
        解析文档并分块
        返回: [{"content": "...", "chunk_index": 0, "metadata": {...}}, ...]
        """
        try:
            logger.info(f"开始解析文档: {file_url}, 类型: {file_type}")

            # 1. 下载文件
            if file_type != 'URL':
                content = self._download_file(file_url)
            else:
                content = None

            # 2. 根据类型解析
            if file_type == 'PDF':
                text = self._parse_pdf(content)
            elif file_type == 'MARKDOWN':
                text = self._parse_markdown(content)
            elif file_type == 'WORD':
                text = self._parse_word(content)
            elif file_type == 'EXCEL':
                text = self._parse_excel(content)
            elif file_type == 'URL':
                text = self._parse_url(file_url)
            else:
                # 默认按文本处理
                text = content.decode('utf-8', errors='ignore')

            # 3. 清理文本
            text = self._clean_text(text)

            if not text or len(text.strip()) < 10:
                raise ValueError("文档内容为空或过短")

            logger.info(f"文档解析完成，文本长度: {len(text)}")

            # 4. 文本分块
            chunks = _split_text(text)

            # 5. 构建结果
            result = []
            for i, chunk in enumerate(chunks):
                result.append({
                    "content": chunk,
                    "chunk_index": i,
                    "metadata": {
                        "file_type": file_type,
                        "char_count": len(chunk),
                        "file_url": file_url
                    }
                })

            logger.info(f"文档分块完成，共 {len(result)} 个切片")
            return result

        except Exception as e:
            logger.error(f"文档解析失败: {str(e)}", exc_info=True)
            raise

    def _download_file(self, url: str) -> bytes:
        """下载文件"""
        try:
            logger.info(f"下载文件: {url}")
            response = requests.get(url, timeout=self.timeout)
            response.raise_for_status()
            logger.info(f"文件下载完成，大小: {len(response.content)} 字节")
            return response.content
        except Exception as e:
            logger.error(f"文件下载失败: {str(e)}")
            raise

    def _parse_pdf(self, content: bytes) -> str:
        """解析PDF"""
        try:
            pdf_file = io.BytesIO(content)
            reader = PyPDF2.PdfReader(pdf_file)

            text = ""
            total_pages = len(reader.pages)
            logger.info(f"PDF共 {total_pages} 页")

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"

                if (i + 1) % 10 == 0:
                    logger.info(f"已处理 {i + 1}/{total_pages} 页")

            logger.info(f"PDF解析完成，提取文本长度: {len(text)}")
            return text

        except Exception as e:
            logger.error(f"PDF解析失败: {str(e)}")
            raise

    def _parse_markdown(self, content: bytes) -> str:
        """解析Markdown"""
        try:
            md_text = content.decode('utf-8', errors='ignore')

            # 转换为HTML
            html = markdown.markdown(md_text)

            # 提取纯文本
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()

            logger.info(f"Markdown解析完成，文本长度: {len(text)}")
            return text

        except Exception as e:
            logger.error(f"Markdown解析失败: {str(e)}")
            raise

    def _parse_word(self, content: bytes) -> str:
        """解析Word文档"""
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)

            # 提取段落文本
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n\n".join(paragraphs)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += "\n" + row_text

            logger.info(f"Word文档解析完成，文本长度: {len(text)}")
            return text

        except Exception as e:
            logger.error(f"Word文档解析失败: {str(e)}")
            raise

    def _parse_excel(self, content: bytes) -> str:
        """解析Excel"""
        try:
            import pandas as pd

            excel_file = io.BytesIO(content)
            df_dict = pd.read_excel(excel_file, sheet_name=None)

            text = ""
            for sheet_name, sheet_df in df_dict.items():
                text += f"\n\n=== {sheet_name} ===\n\n"
                text += sheet_df.to_string(index=False)

            logger.info(f"Excel解析完成，文本长度: {len(text)}")
            return text

        except Exception as e:
            logger.error(f"Excel解析失败: {str(e)}")
            raise

    def _parse_url(self, url: str) -> str:
        """抓取网页内容"""
        try:
            logger.info(f"抓取网页: {url}")
            response = requests.get(url, timeout=self.timeout)
            response.raise_for_status()

            soup = BeautifulSoup(response.content, 'html.parser')

            # 移除script和style标签
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()

            # 提取主要内容
            # 优先查找article、main等标签
            main_content = soup.find('article') or soup.find('main') or soup.find('body')

            if main_content:
                text = main_content.get_text(separator="\n", strip=True)
            else:
                text = soup.get_text(separator="\n", strip=True)

            logger.info(f"网页抓取完成，文本长度: {len(text)}")
            return text

        except Exception as e:
            logger.error(f"网页抓取失败: {str(e)}")
            raise

    def _clean_text(self, text: str) -> str:
        """清理文本"""
        # 移除多余的空白字符
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)

        # 移除特殊字符
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)

        # 去除首尾空白
        text = text.strip()

        return text

    def get_file_info(self, file_url: str, file_type: str) -> Dict:
        """获取文件信息（不解析内容）"""
        try:
            if file_type == 'URL':
                response = requests.head(file_url, timeout=10)
            else:
                response = requests.head(file_url, timeout=10)

            return {
                "url": file_url,
                "type": file_type,
                "size": response.headers.get('Content-Length', 'unknown'),
                "content_type": response.headers.get('Content-Type', 'unknown')
            }
        except Exception as e:
            logger.error(f"获取文件信息失败: {str(e)}")
            return {}

    def validate_file(self, file_url: str, file_type: str, max_size_mb: int = 50) -> bool:
        """验证文件是否可以处理"""
        try:
            info = self.get_file_info(file_url, file_type)

            # 检查文件大小
            if info.get('size') != 'unknown':
                size_mb = int(info['size']) / (1024 * 1024)
                if size_mb > max_size_mb:
                    logger.warning(f"文件过大: {size_mb:.2f}MB > {max_size_mb}MB")
                    return False

            return True

        except Exception as e:
            logger.error(f"文件验证失败: {str(e)}")
            return False
